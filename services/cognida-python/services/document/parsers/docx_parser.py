"""Word 文档解析器。"""

from typing import Any

from .base import BaseParser, ParseResult


class DocxParser(BaseParser):
    """Word (docx) 文档解析器。"""

    @property
    def supported_formats(self) -> list[str]:
        return ["docx", "doc"]

    async def parse(
        self,
        source: str | bytes,
        include_metadata: bool = True,
        extract_tables: bool = False,
        extract_images: bool = False,
        **kwargs: Any,
    ) -> ParseResult:
        """解析 Word 文档。

        Args:
            source: 文件路径或内容
            include_metadata: 是否包含元数据
            extract_tables: 是否提取表格
            extract_images: 是否提取图片
            **kwargs: 其他选项

        Returns:
            解析结果
        """
        try:
            from docx import Document
            from io import BytesIO

            if isinstance(source, str):
                doc = Document(source)
            else:
                doc = Document(BytesIO(source))

            # 提取文本
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            text = "\n\n".join(text_parts)

            # 构建元数据
            metadata = {}
            if include_metadata:
                core_props = doc.core_properties
                metadata = {
                    "format": "docx",
                    "page_count": len(doc.paragraphs),
                    "char_count": len(text),
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "created": str(core_props.created) if core_props.created else "",
                }

            # 提取表格
            tables = None
            if extract_tables:
                tables = self._extract_tables(doc)

            # 提取图片
            images = None
            if extract_images:
                images = self._extract_images(doc)

            return ParseResult(
                success=True,
                text=text,
                metadata=metadata,
                tables=tables,
                images=images,
            )

        except Exception as e:
            return ParseResult(
                success=False,
                text="",
                metadata={},
                error=str(e),
            )

    def _extract_tables(self, doc: Any) -> list[dict[str, Any]]:
        """提取表格。

        Args:
            doc: Document 对象

        Returns:
            表格列表
        """
        tables = []
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)

            # 转换为 Markdown
            markdown = self._table_to_markdown(rows)

            tables.append({
                "table_index": table_idx,
                "rows": rows,
                "markdown": markdown,
            })

        return tables

    def _extract_images(self, doc: Any) -> list[dict[str, Any]]:
        """提取图片信息。

        Args:
            doc: Document 对象

        Returns:
            图片信息列表
        """
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append({
                    "path": rel.target_ref,
                })
        return images

    def _table_to_markdown(self, rows: list[list[str]]) -> str:
        """将表格转换为 Markdown 格式。

        Args:
            rows: 表格行数据

        Returns:
            Markdown 字符串
        """
        if not rows:
            return ""

        lines = []
        for i, row in enumerate(rows):
            cells = [cell.replace("|", "\\|").strip() for cell in row]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                # 添加表头分隔符
                lines.append("|" + "|".join(["---"] * len(cells)) + "|")

        return "\n".join(lines)
